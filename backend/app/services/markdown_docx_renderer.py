"""Markdown-to-DOCX renderer using markdown-it-py and python-docx.

REQ-025 §5.3, §5.5: Parses markdown into AST tokens via markdown-it-py,
then maps tokens to python-docx elements for DOCX generation.

Public API:
- render_docx(markdown_content: str) -> bytes
"""

import io
from urllib.parse import urlparse

from docx import Document as create_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from markdown_it import MarkdownIt

# =============================================================================
# Constants
# =============================================================================

_SAFE_URL_SCHEMES = frozenset({"http", "https", "mailto"})

_MAX_MARKDOWN_LENGTH = 100_000  # ~100KB, sufficient for a resume

# =============================================================================
# URL Safety
# =============================================================================


def _is_safe_url(url: str) -> bool:
    """Check if a URL uses a safe protocol.

    Blocks javascript:, data:, and other dangerous schemes.

    Args:
        url: URL string to validate.

    Returns:
        True if the URL scheme is safe for linking.
    """
    if not url:
        return False
    # Allow fragment-only URLs (anchor links)
    if url.startswith("#"):
        return True
    parsed = urlparse(url)
    return parsed.scheme in _SAFE_URL_SCHEMES


# =============================================================================
# Inline Token Processing
# =============================================================================


def _add_inline_tokens_to_paragraph(paragraph, tokens: list) -> None:
    """Add inline markdown-it tokens as runs to a python-docx paragraph.

    Handles: text, bold, italic, links, code_inline, softbreak.

    Args:
        paragraph: python-docx Paragraph to add runs to.
        tokens: List of inline child tokens from markdown-it.
    """
    bold = False
    italic = False
    link_href: str | None = None

    for token in tokens:
        if token.type == "text":
            run = paragraph.add_run(token.content)
            run.bold = bold
            run.italic = italic
            if link_href:
                _add_hyperlink_to_run(paragraph, run, link_href)
        elif token.type == "code_inline":
            run = paragraph.add_run(token.content)
            run.bold = bold
            run.italic = italic
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        elif token.type == "strong_open":
            bold = True
        elif token.type == "strong_close":
            bold = False
        elif token.type == "em_open":
            italic = True
        elif token.type == "em_close":
            italic = False
        elif token.type == "link_open":
            href = token.attrGet("href") or ""
            if _is_safe_url(href):
                link_href = href
        elif token.type == "link_close":
            link_href = None
        elif token.type in ("softbreak", "hardbreak"):
            paragraph.add_run("\n")


def _add_hyperlink_to_run(paragraph, run, url: str) -> None:
    """Convert a run into a hyperlink within its paragraph.

    Args:
        paragraph: The parent paragraph.
        run: The run element to make into a hyperlink.
        url: The URL to link to.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Style the run as a hyperlink (blue, underlined)
    run.font.color.rgb = None  # Will inherit from hyperlink style
    run.font.underline = True

    # Move the run element inside the hyperlink element
    run_element = run._element
    run_element.getparent().remove(run_element)
    hyperlink.append(run_element)

    paragraph._element.append(hyperlink)


# =============================================================================
# Horizontal Rule
# =============================================================================


def _add_horizontal_rule(document: Document) -> None:
    """Add a horizontal rule (bottom border on an empty paragraph).

    Args:
        document: The python-docx Document to add the rule to.
    """
    paragraph = document.add_paragraph()
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# =============================================================================
# Block Token → DOCX Element Conversion
# =============================================================================


def _tokens_to_docx(tokens: list, document: Document) -> bool:
    """Convert markdown-it block tokens to python-docx elements.

    Args:
        tokens: List of block-level tokens from markdown-it.
        document: python-docx Document to populate.

    Returns:
        True if any elements were added to the document.
    """
    added_elements = False
    i = 0
    ordered_counter = 0

    while i < len(tokens):
        token = tokens[i]

        # --- Headings ---
        if token.type == "heading_open":
            level = int(token.tag[1])  # h1→1, h2→2, etc.
            level = min(level, 4)
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                heading = document.add_heading(level=level)
                heading.clear()
                _add_inline_tokens_to_paragraph(heading, tokens[i + 1].children or [])
                if level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                added_elements = True
                i += 3  # heading_open, inline, heading_close
                continue

        # --- Paragraphs ---
        elif token.type == "paragraph_open":
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                para = document.add_paragraph()
                _add_inline_tokens_to_paragraph(para, tokens[i + 1].children or [])
                added_elements = True
                i += 3  # paragraph_open, inline, paragraph_close
                continue

        # --- Horizontal Rule ---
        elif token.type == "hr":
            _add_horizontal_rule(document)
            added_elements = True
            i += 1
            continue

        # --- Unordered List ---
        elif token.type in ("bullet_list_open", "bullet_list_close"):
            i += 1
            continue

        # --- Ordered List ---
        elif token.type in ("ordered_list_open", "ordered_list_close"):
            ordered_counter = 0
            i += 1
            continue

        # --- List Items ---
        elif token.type == "list_item_open":
            is_ordered = _is_inside_ordered_list(tokens, i)
            if is_ordered:
                ordered_counter += 1

            j = i + 1
            while j < len(tokens) and tokens[j].type != "list_item_close":
                if tokens[j].type == "inline":
                    style = "List Number" if is_ordered else "List Bullet"
                    para = document.add_paragraph(style=style)
                    _add_inline_tokens_to_paragraph(para, tokens[j].children or [])
                    added_elements = True
                j += 1
            i = j + 1
            continue

        # --- Blockquote ---
        elif token.type == "blockquote_open":
            j = i + 1
            bq_tokens: list = []
            depth = 1
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                elif tokens[j].type == "inline":
                    bq_tokens.extend(tokens[j].children or [])
                j += 1
            if bq_tokens:
                para = document.add_paragraph()
                para.paragraph_format.left_indent = Pt(36)
                _add_inline_tokens_to_paragraph(para, bq_tokens)
                # Set italic for blockquote
                for run in para.runs:
                    run.italic = True
                added_elements = True
            i = j
            continue

        # --- Code Block / Fence ---
        elif token.type in ("fence", "code_block"):
            text = token.content.rstrip("\n")
            para = document.add_paragraph(text)
            for run in para.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8)
            added_elements = True
            i += 1
            continue

        i += 1

    return added_elements


def _is_inside_ordered_list(tokens: list, current_idx: int) -> bool:
    """Check if a list_item_open at current_idx is inside an ordered list.

    Walks backwards to find the nearest list open token.

    Args:
        tokens: Full token list.
        current_idx: Index of the list_item_open token.

    Returns:
        True if inside an ordered_list, False otherwise.
    """
    depth = 0
    for j in range(current_idx - 1, -1, -1):
        if tokens[j].type == "list_item_close":
            depth += 1
        elif tokens[j].type == "list_item_open":
            if depth > 0:
                depth -= 1
            else:
                continue
        elif tokens[j].type == "ordered_list_open":
            return True
        elif tokens[j].type == "bullet_list_open":
            return False
    return False


# =============================================================================
# Public API
# =============================================================================


def render_docx(markdown_content: str) -> bytes:
    """Render markdown content to DOCX bytes.

    REQ-025 §5.3: Parses markdown via markdown-it-py, maps tokens to
    python-docx elements, and builds a DOCX document.

    Args:
        markdown_content: Markdown string to render.

    Returns:
        DOCX document as bytes.

    Raises:
        ValueError: If markdown_content is empty or whitespace-only.
    """
    if not markdown_content or not markdown_content.strip():
        raise ValueError("Markdown content cannot be empty.")

    if len(markdown_content) > _MAX_MARKDOWN_LENGTH:
        raise ValueError(
            f"Markdown content exceeds maximum length of {_MAX_MARKDOWN_LENGTH} characters."
        )

    md = MarkdownIt("commonmark")
    tokens = md.parse(markdown_content)

    document = create_document()
    has_content = _tokens_to_docx(tokens, document)

    if not has_content:
        raise ValueError("Markdown content produced no renderable elements.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
