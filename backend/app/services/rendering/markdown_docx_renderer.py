"""Markdown-to-DOCX renderer using markdown-it-py and python-docx.

REQ-025 §5.3, §5.5: Parses markdown into AST tokens via markdown-it-py,
then maps tokens to python-docx elements for DOCX generation.

Public API:
- render_docx(markdown_content: str) -> bytes

Called by: app/api/v1/job_variants.py, app/api/v1/base_resumes.py, and unit tests.
"""

import io
from urllib.parse import urlparse

from docx import Document as create_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from markdown_it import MarkdownIt

# =============================================================================
# Constants
# =============================================================================

_SAFE_URL_SCHEMES = frozenset({"http", "https", "mailto"})

_MAX_MARKDOWN_LENGTH = 100_000  # ~100KB, sufficient for a resume

# =============================================================================
# URL Safety
# =============================================================================


def _is_safe_url(url: str) -> bool:
    """Check if a URL uses a safe protocol.

    Blocks javascript:, data:, and other dangerous schemes.

    Args:
        url: URL string to validate.

    Returns:
        True if the URL scheme is safe for linking.
    """
    if not url:
        return False
    # Allow fragment-only URLs (anchor links)
    if url.startswith("#"):
        return True
    parsed = urlparse(url)
    return parsed.scheme in _SAFE_URL_SCHEMES


def _resolve_safe_href(token) -> str | None:
    """Extract href from a link_open token if the URL scheme is safe.

    Args:
        token: A markdown-it link_open token.

    Returns:
        The href string if safe, None otherwise.
    """
    href = token.attrGet("href") or ""
    return href if _is_safe_url(href) else None


# =============================================================================
# Inline Token Processing
# =============================================================================


def _add_inline_tokens_to_paragraph(paragraph, tokens: list) -> None:
    """Add inline markdown-it tokens as runs to a python-docx paragraph.

    Handles: text, bold, italic, links, code_inline, softbreak.

    Args:
        paragraph: python-docx Paragraph to add runs to.
        tokens: List of inline child tokens from markdown-it.
    """
    bold = False
    italic = False
    link_href: str | None = None

    for token in tokens:
        if token.type == "text":
            run = paragraph.add_run(token.content)
            run.bold = bold
            run.italic = italic
            if link_href:
                _add_hyperlink_to_run(paragraph, run, link_href)
        elif token.type == "code_inline":
            run = paragraph.add_run(token.content)
            run.bold = bold
            run.italic = italic
            run.font.name = "Courier New"
            run.font.size = Pt(8)
        elif token.type == "strong_open":
            bold = True
        elif token.type == "strong_close":
            bold = False
        elif token.type == "em_open":
            italic = True
        elif token.type == "em_close":
            italic = False
        elif token.type == "link_open":
            link_href = _resolve_safe_href(token)
        elif token.type == "link_close":
            link_href = None
        elif token.type in ("softbreak", "hardbreak"):
            paragraph.add_run("\n")


def _add_hyperlink_to_run(paragraph, run, url: str) -> None:
    """Convert a run into a hyperlink within its paragraph.

    Args:
        paragraph: The parent paragraph.
        run: The run element to make into a hyperlink.
        url: The URL to link to.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",  # nosonar — OOXML namespace URI (ECMA-376), not a network URL
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Style the run as a hyperlink (blue, underlined)
    run.font.color.rgb = None  # Will inherit from hyperlink style
    run.font.underline = True

    # Move the run element inside the hyperlink element
    run_element = run._element
    run_element.getparent().remove(run_element)
    hyperlink.append(run_element)

    paragraph._element.append(hyperlink)


# =============================================================================
# Horizontal Rule
# =============================================================================


def _add_horizontal_rule(document: Document) -> None:
    """Add a horizontal rule (bottom border on an empty paragraph).

    Args:
        document: The python-docx Document to add the rule to.
    """
    paragraph = document.add_paragraph()
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# =============================================================================
# Block Token Handlers
# =============================================================================


def _handle_heading_docx(
    tokens: list,
    i: int,
    document: Document,
) -> tuple[int, bool] | None:
    """Handle a heading_open token: create heading with inline content."""
    if i + 1 >= len(tokens) or tokens[i + 1].type != "inline":
        return None
    level = min(int(tokens[i].tag[1]), 4)
    heading = document.add_heading(level=level)
    heading.clear()
    _add_inline_tokens_to_paragraph(heading, tokens[i + 1].children or [])
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return i + 3, True


def _handle_paragraph_docx(
    tokens: list,
    i: int,
    document: Document,
) -> tuple[int, bool] | None:
    """Handle a paragraph_open token: create paragraph with inline content."""
    if i + 1 >= len(tokens) or tokens[i + 1].type != "inline":
        return None
    para = document.add_paragraph()
    _add_inline_tokens_to_paragraph(para, tokens[i + 1].children or [])
    return i + 3, True


def _handle_list_item_docx(
    tokens: list,
    i: int,
    document: Document,
) -> tuple[int, bool]:
    """Handle a list_item_open token: create styled list paragraph(s)."""
    is_ordered = _is_inside_ordered_list(tokens, i)
    added = False
    j = i + 1
    while j < len(tokens) and tokens[j].type != "list_item_close":
        if tokens[j].type == "inline":
            style = "List Number" if is_ordered else "List Bullet"
            para = document.add_paragraph(style=style)
            _add_inline_tokens_to_paragraph(para, tokens[j].children or [])
            added = True
        j += 1
    return j + 1, added


def _handle_blockquote_docx(
    tokens: list,
    i: int,
    document: Document,
) -> tuple[int, bool]:
    """Handle a blockquote_open token: collect content into italic paragraph."""
    j = i + 1
    bq_tokens: list = []
    depth = 1
    while j < len(tokens) and depth > 0:
        if tokens[j].type == "blockquote_open":
            depth += 1
        elif tokens[j].type == "blockquote_close":
            depth -= 1
        elif tokens[j].type == "inline":
            bq_tokens.extend(tokens[j].children or [])
        j += 1
    if not bq_tokens:
        return j, False
    para = document.add_paragraph()
    para.paragraph_format.left_indent = Pt(36)
    _add_inline_tokens_to_paragraph(para, bq_tokens)
    for run in para.runs:
        run.italic = True
    return j, True


def _handle_code_block_docx(
    tokens: list,
    i: int,
    document: Document,
) -> tuple[int, bool]:
    """Handle a fence or code_block token: create monospaced paragraph."""
    text = tokens[i].content.rstrip("\n")
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(8)
    return i + 1, True


def _dispatch_block_docx(
    tokens: list,
    i: int,
    document: Document,
) -> tuple[int, bool] | None:
    """Route a block token to its handler.

    Returns:
        (next_index, element_added) if handled, None otherwise.
    """
    tok = tokens[i]
    if tok.type == "heading_open":
        return _handle_heading_docx(tokens, i, document)
    if tok.type == "paragraph_open":
        return _handle_paragraph_docx(tokens, i, document)
    if tok.type == "hr":
        _add_horizontal_rule(document)
        return i + 1, True
    if tok.type in (
        "bullet_list_open",
        "bullet_list_close",
        "ordered_list_open",
        "ordered_list_close",
    ):
        return i + 1, False
    if tok.type == "list_item_open":
        return _handle_list_item_docx(tokens, i, document)
    if tok.type == "blockquote_open":
        return _handle_blockquote_docx(tokens, i, document)
    if tok.type in ("fence", "code_block"):
        return _handle_code_block_docx(tokens, i, document)
    return None


# =============================================================================
# Block Token → DOCX Element Conversion
# =============================================================================


def _tokens_to_docx(tokens: list, document: Document) -> bool:
    """Convert markdown-it block tokens to python-docx elements.

    Args:
        tokens: List of block-level tokens from markdown-it.
        document: python-docx Document to populate.

    Returns:
        True if any elements were added to the document.
    """
    added_elements = False
    i = 0
    while i < len(tokens):
        result = _dispatch_block_docx(tokens, i, document)
        if result is not None:
            i, added = result
            added_elements = added_elements or added
        else:
            i += 1
    return added_elements


def _is_inside_ordered_list(tokens: list, current_idx: int) -> bool:
    """Check if a list_item_open at current_idx is inside an ordered list.

    Walks backwards to find the nearest list open token.

    Args:
        tokens: Full token list.
        current_idx: Index of the list_item_open token.

    Returns:
        True if inside an ordered_list, False otherwise.
    """
    depth = 0
    for j in range(current_idx - 1, -1, -1):
        if tokens[j].type == "list_item_close":
            depth += 1
        elif tokens[j].type == "list_item_open":
            if depth > 0:
                depth -= 1
            else:
                continue
        elif tokens[j].type == "ordered_list_open":
            return True
        elif tokens[j].type == "bullet_list_open":
            return False
    return False


# =============================================================================
# Public API
# =============================================================================


def render_docx(markdown_content: str) -> bytes:
    """Render markdown content to DOCX bytes.

    REQ-025 §5.3: Parses markdown via markdown-it-py, maps tokens to
    python-docx elements, and builds a DOCX document.

    Args:
        markdown_content: Markdown string to render.

    Returns:
        DOCX document as bytes.

    Raises:
        ValueError: If markdown_content is empty or whitespace-only.
    """
    if not markdown_content or not markdown_content.strip():
        raise ValueError("Markdown content cannot be empty.")

    if len(markdown_content) > _MAX_MARKDOWN_LENGTH:
        raise ValueError(
            f"Markdown content exceeds maximum length of {_MAX_MARKDOWN_LENGTH} characters."
        )

    md = MarkdownIt("commonmark")
    tokens = md.parse(markdown_content)

    document = create_document()
    has_content = _tokens_to_docx(tokens, document)

    if not has_content:
        raise ValueError("Markdown content produced no renderable elements.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
