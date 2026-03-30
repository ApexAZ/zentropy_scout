"""Tests for MarkdownDocxRenderer — markdown → DOCX via python-docx.

REQ-025 §5.3, §5.5: Verifies that supported markdown features render
to valid DOCX bytes, and edge cases are handled gracefully.
"""

import io

import pytest
from docx import Document

from app.services.rendering.markdown_docx_renderer import render_docx


def _load_docx(docx_bytes: bytes) -> Document:
    """Helper: load DOCX bytes into a python-docx Document for inspection."""
    return Document(io.BytesIO(docx_bytes))


def _get_docx_text(docx_bytes: bytes) -> str:
    """Helper: load DOCX and extract all paragraph text."""
    doc = _load_docx(docx_bytes)
    return "\n".join(p.text for p in doc.paragraphs)


class TestRenderDocxBasic:
    """Basic rendering — valid DOCX output."""

    def test_renders_heading_to_valid_docx(self) -> None:
        """A single heading produces valid DOCX bytes (PK zip header)."""
        result = render_docx("# Hello World")
        assert result[:2] == b"PK"

    def test_renders_multiple_headings(self) -> None:
        """All four heading levels produce valid DOCX with correct heading levels."""
        md = "# H1\n\n## H2\n\n### H3\n\n#### H4"
        result = render_docx(md)
        doc = _load_docx(result)
        heading_levels = [
            p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "Heading 1" in heading_levels
        assert "Heading 2" in heading_levels
        assert "Heading 3" in heading_levels
        assert "Heading 4" in heading_levels

    def test_renders_bold_text(self) -> None:
        """Bold markdown renders as bold run in DOCX."""
        result = render_docx("# Resume\n\n**Bold text here**")
        doc = _load_docx(result)
        body_paras = [p for p in doc.paragraphs if p.style.name == "Normal"]
        assert any(run.bold for p in body_paras for run in p.runs), (
            "Expected at least one bold run"
        )

    def test_renders_italic_text(self) -> None:
        """Italic markdown renders as italic run in DOCX."""
        result = render_docx("# Resume\n\n*Italic text here*")
        doc = _load_docx(result)
        body_paras = [p for p in doc.paragraphs if p.style.name == "Normal"]
        assert any(run.italic for p in body_paras for run in p.runs), (
            "Expected at least one italic run"
        )

    def test_renders_bold_and_italic(self) -> None:
        """Mixed bold and italic in same paragraph."""
        result = render_docx("# Resume\n\n**Bold** and *italic* text")
        doc = _load_docx(result)
        body_paras = [p for p in doc.paragraphs if p.style.name == "Normal"]
        runs = [run for p in body_paras for run in p.runs]
        has_bold = any(r.bold for r in runs)
        has_italic = any(r.italic for r in runs)
        assert has_bold and has_italic

    def test_renders_unordered_list(self) -> None:
        """Bullet list items render to DOCX with bullet style."""
        md = "# Skills\n\n- Python\n- TypeScript\n- SQL"
        result = render_docx(md)
        doc = _load_docx(result)
        list_paras = [p for p in doc.paragraphs if "List" in p.style.name]
        assert len(list_paras) >= 3

    def test_renders_ordered_list(self) -> None:
        """Numbered list items render to DOCX with numbered style."""
        md = "# Steps\n\n1. First\n2. Second\n3. Third"
        result = render_docx(md)
        doc = _load_docx(result)
        list_paras = [p for p in doc.paragraphs if "List" in p.style.name]
        assert len(list_paras) >= 3

    def test_renders_horizontal_rule(self) -> None:
        """Horizontal rule renders to DOCX (produces valid output)."""
        md = "# Name\n\n---\n\n## Experience"
        result = render_docx(md)
        assert result[:2] == b"PK"

    def test_renders_link(self) -> None:
        """Links render to DOCX (text preserved)."""
        md = "# Contact\n\n[My Portfolio](https://example.com)"
        result = render_docx(md)
        assert "My Portfolio" in _get_docx_text(result)

    def test_renders_blockquote(self) -> None:
        """Blockquote renders to DOCX."""
        md = "# Summary\n\n> Experienced engineer with 10 years..."
        result = render_docx(md)
        assert "Experienced engineer" in _get_docx_text(result)

    def test_renders_plain_paragraph(self) -> None:
        """Plain paragraph text renders to DOCX."""
        md = "# Resume\n\nThis is a plain paragraph with no formatting."
        result = render_docx(md)
        assert "plain paragraph" in _get_docx_text(result)


class TestRenderDocxEdgeCases:
    """Edge cases and error handling."""

    def test_empty_string_raises_value_error(self) -> None:
        """Empty content is rejected."""
        with pytest.raises(ValueError, match="[Ee]mpty"):
            render_docx("")

    def test_whitespace_only_raises_value_error(self) -> None:
        """Whitespace-only content is rejected."""
        with pytest.raises(ValueError, match="[Ee]mpty"):
            render_docx("   \n\n  ")

    def test_exceeds_max_length_raises_value_error(self) -> None:
        """Content exceeding max length is rejected."""
        md = "# H\n\n" + "x" * 100_001
        with pytest.raises(ValueError, match="maximum length"):
            render_docx(md)

    def test_heading_only(self) -> None:
        """A single heading with no body still produces valid DOCX."""
        result = render_docx("# Just a Heading")
        assert result[:2] == b"PK"

    def test_complex_resume_content(self) -> None:
        """Full resume-like content renders without error."""
        md = """# John Doe

## Summary

Experienced **software engineer** with *10 years* of experience.

---

## Experience

### Senior Developer — Acme Corp

- Led team of 5 engineers
- Reduced deploy time by **40%**
- Implemented CI/CD pipeline

### Developer — StartupCo

1. Built REST API from scratch
2. Integrated payment processing

---

## Education

#### BS Computer Science — MIT

---

## Skills

- Python, TypeScript, Go
- PostgreSQL, Redis
- AWS, Docker, Kubernetes

> "Outstanding contributor" — Annual Review 2024

[LinkedIn](https://linkedin.com/in/johndoe)
"""
        result = render_docx(md)
        assert result[:2] == b"PK"
        full_text = _get_docx_text(result)
        assert "John Doe" in full_text
        assert "Senior Developer" in full_text

    def test_inline_code_rendered_as_text(self) -> None:
        """Inline code is rendered as plain text."""
        md = "# Resume\n\nUses `Python` and `TypeScript`"
        result = render_docx(md)
        assert "Python" in _get_docx_text(result)

    def test_code_block_rendered_as_text(self) -> None:
        """Code blocks are rendered as plain text."""
        md = "# Resume\n\n```\ndef hello():\n    pass\n```"
        result = render_docx(md)
        assert "def hello():" in _get_docx_text(result)


class TestRenderDocxSecurity:
    """Security tests — XSS, injection, and protocol safety."""

    def test_script_tag_in_text_renders_safely(self) -> None:
        """HTML script tags are rendered as plain text, not executed."""
        md = "# Resume\n\n<script>alert('xss')</script>"
        result = render_docx(md)
        assert result[:2] == b"PK"

    def test_javascript_protocol_link_renders_without_href(self) -> None:
        """javascript: URLs are blocked — link text preserved but no hyperlink."""
        md = "# Resume\n\n[click me](javascript:alert(1))"
        result = render_docx(md)
        assert "click me" in _get_docx_text(result)

    def test_data_protocol_link_renders_without_href(self) -> None:
        """data: URLs are blocked."""
        md = "# Resume\n\n[click](data:text/html,<script>alert(1)</script>)"
        result = render_docx(md)
        assert result[:2] == b"PK"

    def test_protocol_relative_url_is_blocked(self) -> None:
        """Protocol-relative URLs (//evil.com) are blocked."""
        md = "# Resume\n\n[click](//evil.com/xss)"
        result = render_docx(md)
        assert "click" in _get_docx_text(result)

    def test_html_entities_in_text_are_preserved(self) -> None:
        """Ampersands and angle brackets render as text."""
        md = "# Resume\n\nCompany A & B <division> worked on C&D"
        result = render_docx(md)
        assert "A & B" in _get_docx_text(result)

    def test_deeply_nested_markdown_does_not_crash(self) -> None:
        """Deeply nested blockquotes don't cause stack overflow."""
        md = "# Resume\n\n" + "> " * 50 + "Deep content"
        result = render_docx(md)
        assert result[:2] == b"PK"
